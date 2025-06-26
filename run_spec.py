# -*- coding: utf-8 -*-
import os
import glob
import docx
import argparse
from tqdm import tqdm
import DatabaseProcess
from os.path import exists
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


'''**
*本範例暫不處理副檔名為.doc
*只處理.docx
*'''
class SpecScanner:
    
    @staticmethod
    def get_file_lists(directory):
        doc_files = [] #紀錄副檔名.doc
        docx_files = [] #紀錄副檔名.docx

        for file in glob.glob(os.path.join(directory, '**'), recursive=True):
            if file.endswith('.doc'):
                doc_files.append(file)
            elif file.endswith('.docx'):
                docx_files.append(file)  
        return doc_files, docx_files
    
    @staticmethod
    def get_filename_without_extension(path):
        filename_with_extension = os.path.basename(path)
        filename_without_extension, _ = os.path.splitext(filename_with_extension)
        return filename_without_extension
    
    
    def __init__(self, directory = "./SVN_datasheet"):
        self.directory = directory 
    
    '''**
    * 讀取文件並將表格資訊抽出，組成特定格式的chunk text
    * @params: file_path - 檔案路徑
    * @return: text_and_tables - 檔案特定格式的chunk text
    *'''
    def read_docx_file(self, file_path:str) -> str:
        doc = docx.Document(file_path)
        text_and_tables = ""
        text_and_tables += 'This is the '+ str(self.get_filename_without_extension(file_path)) +' Specification datasheet.'+ "\n"
        
        # Read paragraphs
        for para in doc.paragraphs:
            text_and_tables += para.text + "\n"

        # Read tables
        for table in doc.tables:
            text_and_tables += "Table: \n"
            for row in table.rows:
                row_data = "\t".join(cell.text for cell in row.cells)
                text_and_tables += row_data + "\n"

        return text_and_tables
    
    '''**
    * 建立Langchain Document列表
    * @params: NA
    * @return: List<langchain.docstore.document>
    *'''
    def scan_folder_and_create_document(self) -> list:
        # 獲取兩種不同附檔名的Word (此範例不處理doc_files)
        doc_files, docx_files = self.get_file_lists(self.directory)
        count = 0        
        doc_list = []
        for docx_file in docx_files:
            try:
                count += 1
                text_and_tables = self.read_docx_file(docx_file)
                metadata = {'source': os.path.basename(docx_file)}
                doc_list.append(Document(page_content=text_and_tables, metadata=metadata))
                print(f"Processed {count} files")
            except Exception as e:
                print(f"Error processing {docx_file}: {e}")             
                
        print(doc_list[:10])
        print("total len of docs: ",len(doc_list))
        
        return doc_list
    
    """逐筆插入文件到 ChromaDB"""
    def insert_documents_for_each(self, lcdb_obj, doc_list):
        for document in tqdm(doc_list,
                             total=len(doc_list), 
                             desc="Inserting documents one by one",
                             unit="pcs"):
            try:
                lcdb_obj.insert_data2db([document])  # 逐筆插入，避免觸發存取速率限制
            except Exception as e:
                print(f"Error inserting document: {document.metadata['source']}, Error: {e}")
    

def main():
    
    parser = argparse.ArgumentParser(description='Spec document')
    parser.add_argument('-d', '--dir', type=str, help='chromadb儲存的資料夾位置')
    parser.add_argument('-cn', '--collection_name', type=str, default="test_collection", help='collection name of chroma db')
    parser.add_argument('-s', '--src', type=str, default="./SVN_datasheet", help='Data Sheet 資料夾位置')
    args = parser.parse_args()
    
    if args.dir is not None:
        # check if the directory exists, if not create it
        if not exists(args.dir):
            os.makedirs(args.dir)
        
        # chroma setting
        chroma_login_info = {
            "chromadb_path": args.dir,
            "collection_name": args.collection_name,
            "embedding_function": DatabaseProcess.embedding_function
        }
        
        # Chroma DB object
        lcdb_chroma = DatabaseProcess.LangchainChromaDB(chroma_login_info)
        lcdb_chroma.init_database()
        
        # Spec document object
        print("Start creating documents....")
        scanner = SpecScanner(args.src)
        doc_list = scanner.scan_folder_and_create_document()
        
        # Insert documents one by one to avoid rate limiting
        scanner.insert_documents_for_each(lcdb_chroma, doc_list)
        print("Done")
        
        
if __name__ == "__main__":
    main()